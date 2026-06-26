from pathlib import Path
from tempfile import TemporaryDirectory
import sys

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.services.parser import parse_document


def run() -> None:
    with TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / "table.docx"
        doc = Document()
        doc.add_paragraph("第一章 项目概述。本项目采用统一数据底座。")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "服务范围"
        table.cell(0, 1).text = "运维支持"
        table.cell(1, 0).text = "响应要求"
        table.cell(1, 1).text = "及时响应"
        doc.save(path)

        parsed = parse_document(path, ".docx")
        assert "运维支持" in parsed.text, parsed.text
        assert "及时响应" in parsed.text, parsed.text
        assert parsed.paragraph_count == 3, parsed.paragraph_count
        assert parsed.sentence_count >= 3, parsed.sentence_count
        print({"paragraph_count": parsed.paragraph_count, "sentence_count": parsed.sentence_count})


if __name__ == "__main__":
    run()
