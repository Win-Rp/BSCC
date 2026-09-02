import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


SENTENCE_SPLIT_RE = re.compile(r"[^。！？!?；;]+[。！？!?；;]*")

SCAN_PDF_MESSAGE = "该 PDF 为扫描件或图片型文件（无文字层），V1 版本暂不支持，请上传可复制文字的 PDF 或 Word 文档"
PDF_INVALID_MESSAGE = "PDF 文件无法读取，可能已损坏或加密"

# 错误码 -> 用户可读中文说明。业务代码抛错统一采用 "CODE:说明" 格式。
ERROR_MESSAGES = {
    "VALIDATION_ERROR": "请求参数错误",
    "UNSUPPORTED_FILE_TYPE": "文件类型不支持，仅支持 Word（.docx）或 PDF 文档",
    "SCAN_PDF_NOT_SUPPORTED": SCAN_PDF_MESSAGE,
    "PDF_INVALID": PDF_INVALID_MESSAGE,
    "PARSE_FILE_FAILED": "文件解析失败（文件可能已损坏、加密或格式异常）",
    "PARSE_FAILED": "文件解析失败",
}


@dataclass
class ParsedDocument:
    text: str
    blocks: list[dict[str, Any]]
    metadata: dict[str, Any]
    page_count: int | None
    word_count: int
    sentence_count: int
    paragraph_count: int


def parse_keywords(raw: str | None) -> list[str]:
    if not raw:
        return []
    parts = re.split(r"[,，\n\r]+", raw)
    return [part.strip() for part in parts if part.strip()]


def parse_document(file_path: Path, file_ext: str) -> ParsedDocument:
    ext = file_ext.lower()
    if ext == ".docx":
        return _parse_docx(file_path)
    if ext == ".pdf":
        return _parse_pdf(file_path)
    raise ValueError(f"UNSUPPORTED_FILE_TYPE:{ERROR_MESSAGES['UNSUPPORTED_FILE_TYPE']}")


def assert_readable_pdf(file_path: Path) -> None:
    """上传前探测 PDF 是否支持解析（含文字层）。

    扫描件/图片型 PDF（无可提取文字）抛 SCAN_PDF_NOT_SUPPORTED；
    损坏、加密等无法读取的 PDF 抛 PDF_INVALID。
    """
    try:
        reader = PdfReader(str(file_path))
        for page in reader.pages:
            if (page.extract_text() or "").strip():
                return
    except Exception as exc:
        raise ValueError(f"PDF_INVALID:{PDF_INVALID_MESSAGE}（{type(exc).__name__}）") from exc
    raise ValueError(f"SCAN_PDF_NOT_SUPPORTED:{SCAN_PDF_MESSAGE}")


def write_parsed_files(parsed: ParsedDocument, text_path: Path, json_path: Path) -> None:
    text_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.parent.mkdir(parents=True, exist_ok=True)
    text_path.write_text(parsed.text, encoding="utf-8")
    json_path.write_text(
        json.dumps(
            {
                "text": parsed.text,
                "blocks": parsed.blocks,
                "metadata": parsed.metadata,
                "page_count": parsed.page_count,
                "word_count": parsed.word_count,
                "sentence_count": parsed.sentence_count,
                "paragraph_count": parsed.paragraph_count,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def load_parsed_json(path: str | Path) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _parse_docx(file_path: Path) -> ParsedDocument:
    doc = Document(str(file_path))
    paragraphs = [_normalize_space(p.text) for p in doc.paragraphs if _normalize_space(p.text)]
    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize_space(cell.text) for cell in row.cells if _normalize_space(cell.text)]
            if cells:
                paragraphs.append(" | ".join(cells))
    blocks = _build_blocks(paragraphs)
    props = doc.core_properties
    metadata = {
        "author": props.author,
        "created": props.created.isoformat() if props.created else None,
        "modified": props.modified.isoformat() if props.modified else None,
        "software": "Microsoft Word / DOCX",
        "template_source": props.category,
        "last_modified_by": props.last_modified_by,
    }
    text = "\n".join(paragraphs)
    return ParsedDocument(
        text=text,
        blocks=blocks,
        metadata=metadata,
        page_count=None,
        word_count=_count_chars(text),
        sentence_count=len(blocks),
        paragraph_count=len(paragraphs),
    )


def _parse_pdf(file_path: Path) -> ParsedDocument:
    reader = PdfReader(str(file_path))
    paragraphs: list[str] = []
    page_numbers: list[int] = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        page_text = _merge_pdf_lines(page_text)
        for paragraph in re.split(r"\n{2,}", page_text):
            clean = _normalize_space(paragraph)
            if clean:
                paragraphs.append(clean)
                page_numbers.append(page_index)
    if not paragraphs:
        raise ValueError(f"SCAN_PDF_NOT_SUPPORTED:{SCAN_PDF_MESSAGE}")
    blocks = _build_blocks(paragraphs, page_numbers)
    metadata_raw = reader.metadata or {}
    metadata = {
        "author": _metadata_value(metadata_raw, "/Author"),
        "created": _metadata_value(metadata_raw, "/CreationDate"),
        "modified": _metadata_value(metadata_raw, "/ModDate"),
        "software": _metadata_value(metadata_raw, "/Creator"),
        "template_source": None,
        "producer": _metadata_value(metadata_raw, "/Producer"),
    }
    text = "\n".join(paragraphs)
    return ParsedDocument(
        text=text,
        blocks=blocks,
        metadata=metadata,
        page_count=len(reader.pages),
        word_count=_count_chars(text),
        sentence_count=len(blocks),
        paragraph_count=len(paragraphs),
    )


def _build_blocks(paragraphs: list[str], page_numbers: list[int] | None = None) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for paragraph_index, paragraph in enumerate(paragraphs, start=1):
        sentences = _split_sentences(paragraph)
        page = page_numbers[paragraph_index - 1] if page_numbers else 1
        for sentence_index, sentence in enumerate(sentences, start=1):
            block_id = f"p{paragraph_index}-s{sentence_index}"
            blocks.append(
                {
                    "block_id": block_id,
                    "page": page,
                    "paragraph": paragraph_index,
                    "sentence": sentence_index,
                    "text": sentence,
                    "char_count": _count_chars(sentence),
                }
            )
    return blocks


def _split_sentences(text: str) -> list[str]:
    parts = [_normalize_space(match.group(0)) for match in SENTENCE_SPLIT_RE.finditer(text)]
    parts = [part for part in parts if part]
    if len(parts) <= 1 and len(text) > 120:
        return [_normalize_space(text[i : i + 80]) for i in range(0, len(text), 80) if _normalize_space(text[i : i + 80])]
    return parts or [text]


def _count_chars(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def _normalize_space(text: str) -> str:
    return re.sub(r"[ \t\r\f\v]+", " ", text.replace("\u3000", " ")).strip(" \t\r\f\v")


def _merge_pdf_lines(text: str) -> str:
    lines = [_normalize_space(line) for line in text.splitlines()]
    merged: list[str] = []
    buffer = ""
    for line in lines:
        if not line:
            if buffer:
                merged.append(buffer)
                buffer = ""
            merged.append("")
            continue
        if buffer and not re.search(r"[。！？!?；;：:]$", buffer):
            buffer = f"{buffer}\n{line}"
        else:
            if buffer:
                merged.append(buffer)
            buffer = line
    if buffer:
        merged.append(buffer)
    return "\n\n".join(merged)


def _metadata_value(metadata: Any, key: str) -> str | None:
    value = metadata.get(key)
    return str(value) if value else None
